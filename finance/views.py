from django.shortcuts import render
from django.http import HttpResponseRedirect, JsonResponse, HttpResponse
from django.utils.http import urlencode
from django.urls import reverse
from django.contrib.auth.models import User
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout
from .models import *
from django.db.models import Q, Sum
from django.core.paginator import Paginator
from json import loads, dumps
from urllib.parse import unquote
import math
import datetime


# Functions
def companyDetails():
    data = Settings.objects.first()
    if data :
        return data.serialize()
    return {}

# JSON Requests
def getInfo(request, id=None, name=None):
    kind =  request.build_absolute_uri().split('/')[-2]
    if kind == 'item':
        item = Items.objects.filter(Q(code=id) | Q(pk=id) | Q(item=name))
        if item :
            return JsonResponse({'item' : item[0].serialize()}, status=200)
        
    elif kind == 'category':
        category = ItemsCategories.objects.filter(Q(code=id) | Q(pk=id) | Q(category=name))
        if category :
            return JsonResponse({'category' : category[0].serialize()}, status=200)
    else :
        return JsonResponse({'message' : 'Bad Request!'}, status=404)
    return JsonResponse({'message' : 'المنتج غير موجود'}, status=403)

def getRecommendations(request, table, text):
    text = unquote(text)
    if table == 'user':
        list = [{'value': user.item, 'id': user.pk} for user in Users.objects.filter(name__contains=text)[:5]]
        pass
    elif table == 'inventory':
        list = [{'value': inventory.inventory, 'id': inventory.pk} for inventory in Inventories.objects.filter(inventory__contains=text)[:5]]
    elif table == 'item':
        list = [{'value': item.item, 'id': item.pk} for item in Items.objects.filter(item__contains=text)[:5]]
    elif table == 'category':
        list = [{'value': category.category, 'id': category.pk} for category in ItemsCategories.objects.filter(category__contains=text)[:5]]
    else :
        return JsonResponse({'message' : 'Bad Request!'}, status=404)
    return JsonResponse({'data' : list}, status=200)

@csrf_exempt
def checkAvailability(request):
    data = loads(request.body)
    if data['transactionType'] == 'sell':
        quantity = int(data['quantity'])
        item = Items.objects.get(code=int(data['code']))
        inventory = Inventories.objects.get(pk=int(data['inventory']))
        # Check item quantity
        if quantity > item.quantity :
            return JsonResponse({'message' : f'الكميه المطلوبه لمنتج {item.item} غير متوفره'}, status=400)
        row = inventory.itemsList.filter(item=item)
        if not row or quantity > row[0].quantity :
            return JsonResponse({'message' : f'الكميه المطلوبه لمنتج {item.item} غير متوفره في مخزن {inventory.inventory}'}, status=400)
    return JsonResponse({}, status=200)

def paginateRows(request, model):
    # Get Parameters
    try :
        pageNumber = int(request.GET.get('page', 1))
        rowsNumber = int(request.GET.get('rows', 20))
    except :
        pageNumber = 1
        rowsNumber = 20

    inventoryPaginator = Paginator(model, rowsNumber)
    currentPage = inventoryPaginator.get_page(pageNumber)
    
    return (currentPage, rowsNumber)

def calculateInventory(inventory):
    balance = 0
    for data in inventory.itemsList.all():
        data = data.serializeItems()
        quantityPerInventory = data['quantity']
        pricePerItem = data['item']['sellPrice']
        balance += (math.trunc(pricePerItem * 100) * quantityPerInventory) / 100
    return balance

# Create your views here
def logoutView(request):
    logout(request)
    return HttpResponseRedirect(reverse('loginView'))

def loginView(request):
    if request.method == 'GET' :
        message = request.GET.get('message', '')
        return render(request, 'finance/login.html', {
            'message' : message
        })
    elif request.method == 'POST' :
        username = request.POST["username"]
        password = request.POST["password"]
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return HttpResponseRedirect(reverse('inventory'))
        else:
            message = 'اسم المستخدم او كلمه المرور غير صحيحه برجاء المحاوله مره اخري'
            return HttpResponseRedirect(reverse('loginView') + f'?message={message}')


@login_required(login_url="loginView")
def index(request):
    return HttpResponseRedirect(reverse('inventory'))

@login_required(login_url="loginView")
def inventory(request):
    """
    GET :
        When access page it displays a form that has name, balance and country
        Find a search bar 
        Find a table displays name, balance, country, button to generate report of inventory

    POST :
        Add new inventory details
    
    PUT :
        Check if need information about inventory 
            send it back as JSON
        else
            Update inventory details
    """
    if request.method == 'GET':
        # Get Parameters
        refreshCondition = request.GET.get('refresh', True)
        inventoryName = request.GET.get('inventory', '')
        inventoryID = request.GET.get('id', '')
        krwags = {}

        # Make filters
        if inventoryName :
            krwags['inventory__contains'] = inventoryName
        if inventoryID :
            krwags['pk'] = inventoryID

        # Retrieve Data
        inventories = Inventories.objects.filter(**krwags).order_by('id')

        currentPage, rowsNumber = paginateRows(request, inventories)
        for item in currentPage:
            item.balance = 0
            for quantity in item.itemsList.all():
                item.balance += quantity.quantity
        # inventoryPaginator = Paginator(inventories, rowsNumber)
        # currentPage = inventoryPaginator.get_page(pageNumber)

        # Send Data
        if refreshCondition == True:
            data = request.session.get('dataInventory', {})
            if not data :
                data['kind'] = 'addInventory'
            request.session['dataInventory'] = {}
            return render(request, 'finance/inventory.html',{
                'inventories' : currentPage,
                'rows' : int(rowsNumber),
                'rowsOptions' : [1,3,5,10,20,50,100],
                'name' : inventoryName,
                'id' : inventoryID,
                'data' : data,
                'arName' : companyDetails()['arName']
            })
        else :
            hasPrevious = currentPage.has_previous()
            hasNext = currentPage.has_next()
            currentPage = [inventory.serialize() for inventory in currentPage]
            return JsonResponse({
                'inventories' : currentPage,
                'hasNext' : hasNext,
                'hasPrevious' : hasPrevious,
                }, status=200, safe=False)
    
    elif request.method == 'POST':
        name = request.POST['name'] 
        place = request.POST['place']
        balance = request.POST['balance'] or 0
        if name and place:
            Inventories.objects.create(inventory=name, place=place, balance=balance)
        return HttpResponseRedirect(reverse('inventory'))

    elif request.method == 'PUT':
        data =  loads(request.body)
        id =  data['id']
        inventory = Inventories.objects.filter(pk=id)
        if inventory :
            inventory = inventory[0]

            if data['getInfo'] :
                return JsonResponse({'inventory' : inventory.serialize()}, status=200)
            
            if data['update']:
                # Suggest of changing in items and wastes
                name = data['name']
                place = data['place']
                balance = data['balance']
                inventory.inventory = name
                inventory.place = place
                inventory.balance = balance
                inventory.save()
                return HttpResponseRedirect(reverse('inventory'))
                
        return JsonResponse({'message' : 'Not Found!'}, status=404)
        
    pass

def searchItemInventory(request):
    if request.method == 'GET' : 
        data = {}
        data['itemNameSearch'] = item = request.GET.get('itemSearch', '')
        data['inventoryNameSearch'] = inventory = request.GET.get('inventoryName', '')
        inventoryObj = Inventories.objects.filter(inventory=inventory) 
        itemObj = Items.objects.filter(item=item) 
        if itemObj and inventoryObj :
            itemInventory = Inventory_Items.objects.filter(inventory=inventoryObj[0], item=itemObj[0])
            if itemInventory and itemInventory[0].quantity :
                itemInventory = itemInventory[0]
                data['availableSearch'] = 'متوفر'
                data['quantitySearch'] = itemInventory.quantity
            else :
                data['availableSearch'] = "غير متوفر" 
            
        else :
            data['message'] = 'برجاء ادخال كل الاماكن المطلوبه..!'
        data['kind'] = 'checkItem'
        request.session['dataInventory'] = data
        return HttpResponseRedirect(reverse('inventory'))

@login_required(login_url="loginView")
def addItemInventory(request):
    kwrags = {}
    if request.method == 'POST' :
        # Get Data
        kwrags['inventory'] = inventory = request.POST.get('inventoryName' , '')
        kwrags['itemName'] = itemName = request.POST.get('itemAdd', '')
        kwrags['itemCode'] = code = request.POST.get('itemCodeAdd', '')
        kwrags['quantity'] = quantity = request.POST.get('quantityFromLastYear', '')
        try : 
            quantity = int(quantity)
        except :
            quantity = ''
        safeMode = request.POST.get('safeMode', False)
        failed = True
        if itemName or code :
            if quantity and quantity >= 0 :
                if inventory :
                    if not safeMode:
                        try :
                            message = 'المخزن غير موجود في قواعد البيانات'
                            inventoryObj = Inventories.objects.get(inventory=inventory)
                            message = 'البضاعه غير موجوده في قواعد البيانات برجاء ادخال المعلومات الصحيحه'
                            itemObj = Items.objects.get(Q(code=code) | Q(item=itemName))
                            itemInventoryObj = Inventory_Items.objects.filter(inventory=inventoryObj, item=itemObj)
                            if itemInventoryObj :
                                if safeMode :
                                    message = 'يوجد رصيد اول المده من قبل لا يمكنك تحديد اخر'
                                else :
                                    currentQuantity = itemInventoryObj[0].quantityFromLastYear
                                    itemInventoryObj[0].quantityFromLastYear = quantity
                                    itemInventoryObj[0].quantity = itemInventoryObj[0].quantity - currentQuantity + quantity
                                    itemInventoryObj[0].save()
                                    itemObj.quantity = itemObj.quantity - currentQuantity + quantity

                            else :
                                Inventory_Items.objects.create(inventory=inventoryObj, item=itemObj, quantity=quantity, quantityFromLastYear=quantity)
                                itemObj.quantity += quantity

                            inventoryObj.balance = calculateInventory(inventoryObj)
                            # Save Data
                            inventoryObj.save()
                            itemObj.save()
                            failed = False
                            message = 'تمت الاضافه بنجاح'
                        except :
                            pass
                else :
                    message = 'المخزن غير موجود في قواعد البيانات'
            else :
                message = 'برجاء اضافه كميه بالارقام'
        else :
            message = 'الاول البضاعه غير موجوده في قواعد البيانات برجاء ادخال المعلومات الصحيحه' 

        kwrags['message'] = message
        kwrags['failed'] = failed
        kwrags['kind'] = 'addItem'
        request.session['dataInventory'] = kwrags
        return HttpResponseRedirect(reverse('inventory'))
    else :
        return JsonResponse({'message' : 'Bad Request'}, status=404)
    pass

@login_required(login_url="loginView")
def transferItemInventory(request):
    pass

@login_required(login_url="loginView")
def items(request):
    """
        GET : 
            Display form 
                0ne => [Category, Code, ParentCategory] => Types
                Two => [Item, Code, Category, purchasePrice, sellPrice, Quantity, unit, quantityFromLastYear]
            Table display all Types or Items

        POST : 
            Add new Item/Category

        PUT :
            Update Item/Category

        DELETE :
            if items :
                check quantity == 0
                    delete item
                else    
                    return Error
            elif category :
                get all items with this parent
                for each item check quantity == 0
                    delete category and items
                else    
                    return Error

    """
    current_url = request.build_absolute_uri().split('/')

    if request.method == 'GET':
        data = {}
        # response = get_response(request)
        if request.session.get('redirected', False):
            data = request.GET.dict()
        kind = request.GET.get('kind', current_url[-1])
        name = request.GET.get('name', '') 
        code = request.GET.get('code', '') 
        refresh = request.GET.get('refresh', True) 
        kwrags = {}
        itemsRows = []
        categoriesRows = []
        if name :
            if kind == 'items' or kind == 'item':
                kwrags['item__contains'] = name
            else :
                kwrags['category__contains'] = name

        if code :
            kwrags['code'] = code

        rows = 20
        if kind == 'items' or kind == 'item':
            items = Items.objects.filter(**kwrags).order_by('item')
            currentPageItems, rowsPerPageItems = paginateRows(request, items)
            for items in currentPageItems :
                item = items.serialize()
                item['reportLink']  = reverse('itemReport', kwargs={'id' : item['id']})
                itemsRows.append(item)
            rows = rowsPerPageItems
        else : 
            categories = ItemsCategories.objects.filter(**kwrags).order_by('category')
            currentPageCategories, rowsPerPageCategories = paginateRows(request, categories)
            categoriesRows = [category.serialize() for category in currentPageCategories]
            rows = rowsPerPageCategories


        if kind == 'items' or kind == 'item' :
            currentPage = currentPageItems
        else :
            currentPage = currentPageCategories

        if refresh == True:
            return render(request, 'finance/Items.html',{
                'categories' : categoriesRows,
                'items' : itemsRows,
                'rows' : int(rows),
                'rowsOptions' : [1,3,5,10,20,50,100],
                'currentPage' : currentPage,
                'kind' : kind,
                'data' : data,
                'arName' : companyDetails()['arName']
            })
        else :
            if kind == 'items' or kind == 'item' :
                hasPrevious = currentPageItems.has_previous()
                hasNext = currentPageItems.has_next()
            else :
                hasPrevious = currentPageCategories.has_previous()
                hasNext = currentPageCategories.has_next()

            return JsonResponse({
                'categories' : categoriesRows,
                'items' : itemsRows,
                'hasNext' : hasNext,
                'pagesCount' : currentPage.paginator.num_pages,
                'hasPrevious' : hasPrevious,
                }, status=200, safe=False)
    
    elif request.method == 'POST':
        failed = False
        message = ''
        kwrags = {}
        kind = request.POST['kind']
        if request.POST['kind'] == 'items' or request.POST['kind'] == 'item':
            kwrags['item'] = name = request.POST.get('name') 
            kwrags['codeItem'] = code = request.POST.get('code')
            kwrags['categoryItem'] = category = request.POST.get('categoryItem')
            kwrags['purchasePrice'] = purchasePrice = request.POST.get('purchasePrice') or 0
            kwrags['sellPrice'] = sellPrice = request.POST.get('sellPrice') or 0
            kwrags['unit'] = unit = request.POST.get('unit')
            # quantityFromLastYear = request.POST.get('quantityFromLastYear') or 0
            if name and code and category and unit:
                try : 
                    message = 'الصنف الذي تحاول ربطه بالبضاعه به خطأ'
                    categoryObj = ItemsCategories.objects.get(category=category)
                    message = 'بيانات التي تم ادخالها قد تكون مسجله بالفعل الرجاء المحاوله مجددا'
                    Items.objects.create(item=name, code=code, category=categoryObj, unit=unit, sellPrice=sellPrice, purchasePrice=purchasePrice)
                except :
                    failed = True
        else :
            safeMode = request.POST.get('safeMode', False)
            kwrags['category'] = name = request.POST['name'] 
            kwrags['codeCategory'] = code = request.POST['code']
            kwrags['categoryParent'] = category = request.POST.get('category')
            categoryObj = None
            if name and code:
                if category:
                    # Turn on or off safety
                    try: 
                        if safeMode: 
                            message = 'المشتق منه غير موجود برجاء المحاوله مجددا'
                            categoryObj = ItemsCategories.objects.get(category=category)
                        else :
                            try:
                                categoryObj = ItemsCategories.objects.get(category=category)
                            except ItemsCategories.DoesNotExist:
                                categoryObj = None
                    except :
                        failed = True

                if not failed:
                    try :
                        ItemsCategories.objects.create(category=name, code=code, categoryParent=categoryObj)
                    except : 
                        failed = True
                        message = 'بيانات التي تم ادخالها قد تكون مسجله بالفعل الرجاء المحاوله مجددا'

        kwrags['message'] =  message
        kwrags['failed'] = failed
        kwrags['kind'] = kind
        if failed :
            request.session['redirected'] = True
            params = urlencode(kwrags)
            # print(current_url)
            return HttpResponseRedirect(reverse(current_url[-1])+f'?{params}') 
        return HttpResponseRedirect(reverse(current_url[-1])) 


    elif request.method == 'DELETE':
        data = loads(request.body)
        id = int(data['id'])
        kind = data['kind']
        if kind == 'items' or kind == 'item':
            item = Items.objects.get(pk=id)
            if item.quantity == 0:
                item.delete()
                return JsonResponse({'message' : 'تم حذف المنتج بنجاح'}, status=200)

        else:
            category = ItemsCategories.objects.get(pk=id)
            if not any(item.quantity != 0 for item in category.items.all()) :
                for item in category.items.all():
                    # Edit if he need to leave items without deletion.. if he need this syntax then change models easier!
                    item.delete()
                category.delete()
                return JsonResponse({'message' : "تم حذف الصنف بالمنتجات المتعلقه بيه"}, status=200)
        return JsonResponse({'message' : "لا يمكن تنفيذ العمليه لسبب وجود كميه من المنتج او الصنف"}, status=403)
            
    pass

@login_required(login_url="loginView")
def transactions(request):
    """
    GET :
        display Form
            user | items | unit | price | Transactions
        search ['date', 'ID', 'item', 'client']
        Table
    
    POST:
        Add transaction
        if buy :
            update Items => quantity | PurchasePrice | SellPrice
            update Inventory_Items => quantity 
            add ItemPrices record
            add Transaction
            add Transaction_Items
        if sell :
            update Items => quantity 
            update Inventory_Items => quantity 
            add ItemPrices record
            add Transaction
            add Transaction_Items
    """
    if request.method == 'GET':
        # [kind , client, id, fromPrice, toPrice]
        # Get Params
        refreshCondition = request.GET.get('refresh', True)
        type1 = request.GET.get('transactionType', '')
        transactionID = request.GET.get('id', '')
        client = request.GET.get('client', '')
        fromPrice = request.GET.get('fromPrice', '')
        toPrice = request.GET.get('toPrice', '')
        fromDate = request.GET.get('fromDate', '')
        toDate = request.GET.get('toDate', '')
        
        krwags = {}
        # Make filters
        if type1 and type1 != 'all':
            krwags['transactionType'] = type1
            
        if transactionID:
            krwags['pk'] = transactionID

        if client:
            krwags['user__contains'] = client

        if fromPrice:
            krwags['totalPrice__gte'] = fromPrice

        if toPrice:
            krwags['totalPrice__lte'] = toPrice
        
        if fromDate :
            fromDate = datetime.datetime.strptime(fromDate, '%Y-%m-%d')
            krwags['dateTime__gte'] = fromDate
        
        if toDate :
            toDate += ' 23:59:59'
            toDate = datetime.datetime.strptime(toDate, '%Y-%m-%d %H:%M:%S')
            krwags['dateTime__lte'] = toDate

        # Process Data
        inventories = [inventory.serialize() for inventory in Inventories.objects.all()]
        users = [user.name for user in Users.objects.all()]
        items = [item.serialize() for item in Items.objects.all()]
        transactions = Transactions.objects.filter(**krwags).order_by('-dateTime')
        currentPage, rowsNumber = paginateRows(request, transactions)
        transactionsRows = [
            {
                **transaction.serialize(),
                'urlInvoice': reverse('invoice', args=[transaction.serialize()['id']])
            }
            for transaction in currentPage
        ]

        if refreshCondition == True:
            return render(request, 'finance/transactions.html', {
                'users' : users,
                'inventories' : inventories,
                'items' : items,
                'transactions' : transactionsRows,
                'currentPage' : currentPage,
                'rows' : int(rowsNumber),
                'rowsOptions' : [1,3,5,10,20,50,100],
                'arName' : companyDetails()['arName'],
                'taxes' : companyDetails()['taxes'] or 0
            })
        else :
            hasPrevious = currentPage.has_previous()
            hasNext = currentPage.has_next()
            pagesCount = currentPage.paginator.num_pages
            return JsonResponse({
            'transactions' : transactionsRows,
            'hasNext' : hasNext,
            'hasPrevious' : hasPrevious,
            'pagesCount' : pagesCount
            }, status=200, safe=False)

    elif request.method == 'POST':
        data = loads(request.body)
        # print(data)
        transactionType = data['transactionType']
        user = data['user']
        items = data['items']
        totalPrice = math.trunc(float(data['totalPrice']) * 100) / 100
        tax = math.trunc(float(data['tax']) * 100) / 100
        discount = data['discount']
        changePrice = data['changePrice']
        
        if not items:
            return JsonResponse({'message' : 'Access Forbidden!'}, status=403)
        # Change User to user model instead of saving as string [TODO]
        transaction = Transactions.objects.create(transactionType=transactionType, totalPrice=totalPrice, user=user, tax=tax, discount=discount)
        for product in items:
            price = math.trunc(float(product['price']) * 100) / 100
            # inventory of the item
            inventory = Inventories.objects.get(pk=int(product['inventory']))
            # update Items => quantity | PurchasePrice | SellPrice
            item = Items.objects.get(code=product['code'])
            quantity = int(product['quantity'])
            if transactionType == 'purchase':
                item.purchasePrice = price
                if changePrice:
                    item.sellPrice = math.trunc(price * 1.1 * 100) / 100
                item.quantity += quantity
                # update Inventory_Items => quantity 
                row = inventory.itemsList.filter(item=item)
                if row :
                    row[0].quantity += quantity
                else :
                    inventory_item = Inventory_Items.objects.create(item=item, quantity=quantity, inventory=inventory)
                    inventory.itemsList.add(inventory_item)

            elif transactionType == 'sell':
                if changePrice:
                    item.sellPrice = price
                # update quantity of item
                item.quantity -= quantity
                # update Inventory_Items => quantity 
                row = inventory.itemsList.filter(item=item)
                if row :
                    row[0].quantity -= quantity
 
            # add ItemPrices record
            ItemPrices.objects.create(item=item, price=price, transactionType=transactionType)
            # add Transaction_Items
            transaction_item = Transactions_Items.objects.create(item=item, quantity=int(product['quantity']), inventory=inventory, transaction=transaction, price=price)
            transaction.itemsList.add(transaction_item)
            # Save Item and Inventory_Items
            item.save()
            if row : 
                row[0].save()

            # update inventory balance
            inventory.balance = calculateInventory(inventory)
            # Save inventory data
            inventory.save()


        """
                8.5   100 => 850 / 100 => 
                8.5%  100%
        """
        # Save Taxes rate
        settingsData = Settings.objects.first()
        price = (math.trunc(totalPrice * 100) - math.trunc(tax * 100)) / 100
        taxPercent = ((math.trunc(tax * 100) * 100) / (math.trunc(price * 100)))
        if settingsData:
            settingsData.taxes = taxPercent
            settingsData.save()
        else :
            Settings.objects.create(taxes=taxPercent)
        return JsonResponse({'message' : 'تم اضافه الفاتوره بنجاح'}, status=200)

@login_required(login_url="loginView")
def invoice(request, id):
    # Get Item
    transaction = Transactions.objects.get(pk=id).serialize()

    # Edit Values
    if transaction['transactionType'] == 'purchase':
        transaction['transactionType'] = 'المشتريات'
    else:
        transaction['transactionType'] = 'المبيعات'
    # Add derived values
    transaction['inventories'] = set()
    for item in transaction['items']:
        transaction['inventories'].add(item['inventory'])
        item['pricePerItem'] = math.trunc(item['pricePerItem'] * 100) / 100
        item['total'] = math.trunc(item['quantity'] * item['pricePerItem'] * 100) / 100

    transaction['totalBeforeTax'] = (math.trunc(transaction['totalPrice'] * 100) - math.trunc(transaction['tax'] * 100)) / 100

    # pdf = generate_pdf('finance/invoice.html', {'transaction' : transaction})
    # return render_to_pdf_response('finance/invoice.html', {'transaction' : transaction})
    return render(request, 'finance/invoice.html', {'transaction' : transaction, **companyDetails()})

@login_required(login_url="loginView")
def inventoryReport(request, id):
    """
        Get Inventory details [inventory - number of categories - number of items]  
            => [Category with Total Price] 
            => [Item Code Quantity SellPrice TotalPrice]
            {
                'inventory' : ,
                'place' : ,
                'balance' : ,
                'numberOfCategories' : ,
                'numberOfItems' : ,
                'categories' : [{
                    'category' : ,
                    'balance' : ,
                    'items' : {
                        'item' : ,
                        'code' : ,
                        'inventoryItemQuantity': ,
                        ''quantityFromLastYear'': ,
                        'sellPrice': ,
                        'totalPrice' : 
                    }
                }]
            }
    
    """
    inventory = Inventories.objects.get(pk=id)
    data =  inventory.serialize()
    numberOfCategories = 0
    numberOfItems = 0
    categories = {}
    for result in data['items']:
        # Make Item as Dict
        numberOfItems += 1
        totalPriceItem = result['quantity'] * math.trunc(result['item']['sellPrice'] * 100) / 100
        items = {
            'item' : result['item']['item'],
            'code' : result['item']['code'],
            'quantityFromLastYear' : result['quantityFromLastYear'], 
            'quantity' : result['quantity'],
            'sellPrice' : result['item']['sellPrice'],
            'totalPrice' : totalPriceItem
        }

        # Make Category Dict
        category = result['item']['category']
        if category not in categories :
            numberOfCategories += 1
            categories[category] = {
                'category' : category,
                'balance' : totalPriceItem,
                'items' : [items]
            }
        else :
            categories[category]['balance'] += totalPriceItem
            categories[category]['items'].append(items)

    del data['items']

    data['numberOfCategories'] = numberOfCategories
    data['numberOfItems'] = numberOfItems
    data['categories'] = []
    for val in categories.values():
        data['categories'].append(val)
    
    

    return render(request, 'finance/inventoryReport.html', {'data' : data, **companyDetails()})

@login_required(login_url="loginView")
def itemReport(request, id):
    """
        => Item [Item - Code - Total Quantity - Category - Purchase Price - Sell Price - Unit - Number of Inventories] [opt. => Percentage of Total Balance of Category]
        => Inventory [Name - QuantityFromLastYear - CurrentQuantity - TotalPrice]
        => Transactions [
            Overview => Number of Transactions - Number of Purchase Transactions - Number of Sell Transactions 
            - Total Net Profit - Total Purchases - Total Sells 
            2 Tables => TransactionID - Client - Quantity - Price - TotalPrice
        ]
        => ItemPrices OverTime as Table

    """
    # Get Item
    item = Items.objects.get(pk=id)
    data = item.serialize()
    data['totalBalance'] = data['quantity'] * math.trunc(data['sellPrice'] * 100) / 100
    # Add inventoriesNumber to dictionary and balance to every inventory
    data['inventoriesNumber'] = len(data['inventories'])
    for inventory in data['inventories']:
        inventory['balance'] = inventory['quantity'] * math.trunc(data['sellPrice'] * 100) / 100

    # Get Transactions with specific Item
    data['transactions'] = {}
    data['transactions']['totalTransactions'] = 0 
    data['transactions']['purchaseTransactions'] = 0 
    data['transactions']['sellTransactions'] = 0 
    data['transactions']['netProfit'] = 0 
    data['transactions']['totalPurchase'] = 0 
    data['transactions']['totalSell'] = 0 
    data['transactions']['purchaseTable'] = [] 
    data['transactions']['sellTable'] = [] 

    transactions = Transactions.objects.filter(itemsList__in=Transactions_Items.objects.filter(item=item))
    transactionsIDs = []
    for transaction in transactions:
        id = transaction.pk
        if id in transactionsIDs:
            continue
        transactionsIDs.append(id)
        # Total Update
        data['transactions']['totalTransactions'] += 1

        # Assign Variables
        transactionItem = transaction.itemsList.filter(item=item)
        quantity = 0
        price = transactionItem[0].price
        for itemQuantity in transactionItem:
            transactionsIDs.append(itemQuantity.pk)
            quantity += itemQuantity.quantity

        totalPrice = quantity * math.trunc(price * 100) / 100
        transactionTable = {
            'id' : transaction.pk,
            'client' : transaction.user,
            'quantity' : quantity,
            'price' : price,
            'totalPrice' : totalPrice,
            'date' : transaction.dateTime.strftime('%d/%m/%Y')
        }

        # Specify Transaction Type and Update it
        if transaction.transactionType == 'purchase' :
            data['transactions']['totalPurchase'] += 1
            data['transactions']['purchaseTransactions'] += totalPrice
            data['transactions']['netProfit'] -= totalPrice
            data['transactions']['purchaseTable'].append(transactionTable)

        elif transaction.transactionType == 'sell':
            data['transactions']['totalSell'] += 1
            data['transactions']['sellTransactions'] += totalPrice 
            data['transactions']['netProfit'] +=  totalPrice
            data['transactions']['sellTable'].append(transactionTable)
        
    data['transactions']['netProfit'] = math.trunc(data['transactions']['netProfit'] * 100) / 100
    data['transactions']['transactionTypes'] = [
        {
            'purchaseTable' : 'فاتوره المشتريات'
        },
        {
            'sellTable' : 'فاتوره المبيعات'
        }
    ]
    return render(request, 'finance/itemReport.html', {'data' : data, **companyDetails()})
    # return JsonResponse({'data' : data}, status=200)

@login_required(login_url="loginView")
def categoryReport(request, id):
    """
        1 => Get Category details [category - parent of categories - number of children - number of items]  
            2 => [Child Category - has another children [number of children] - Total Price] 
                3 => if Children :
                    4 => Go Step 2   
                else :
                    4 => [Item Code Quantity SellPrice TotalPrice]   
                    5 => [Inventory QuantityFromLastYear CurrentQuantity TotalPrice]
            {
                'category' : ,
                'categoryParent' : ,
                'code' : ,
                'numberOfChildrenCategories' : , [Including sub-children]
                'numberOfItems' : , [Including sub-items]
                'categoryChildren' : [{
                    'category' : ,
                    'totalPrice' : ,
                    'numberOfChildrenCategories' : ,
                    'categoryChildren' : ,
                    'items' : {
                        'item' : ,
                        'code' : ,
                        'inventoryItemQuantity': ,
                        ''quantityFromLastYear'': ,
                        'sellPrice': ,
                        'totalPrice' : 
                    }
                }]
            }
    
    """
    inventory = Inventories.objects.get(pk=id)
    data =  inventory.serialize()
    numberOfCategories = 0
    numberOfItems = 0
    categories = {}
    for result in data['items']:
        # Make Item as Dict
        numberOfItems += 1
        totalPriceItem = result['quantity'] * math.trunc(result['item']['sellPrice'] * 100) / 100
        items = {
            'item' : result['item']['item'],
            'code' : result['item']['code'],
            'quantityFromLastYear' : result['quantityFromLastYear'], 
            'quantity' : result['quantity'],
            'sellPrice' : result['item']['sellPrice'],
            'totalPrice' : totalPriceItem
        }

        # Make Category Dict
        category = result['item']['category']
        if category not in categories :
            numberOfCategories += 1
            categories[category] = {
                'category' : category,
                'balance' : totalPriceItem,
                'items' : [items]
            }
        else :
            categories[category]['balance'] += totalPriceItem
            categories[category]['items'].append(items)

    del data['items']

    data['numberOfCategories'] = numberOfCategories
    data['numberOfItems'] = numberOfItems
    data['categories'] = []
    for val in categories.values():
        data['categories'].append(val)
    

    return render(request, 'finance/inventoryReport.html', {'data' : data})


def reportPage(request):
    """
    [Inventory || Items] 
    Inventories, From/To Date, From/To Item
    (Name - Code - Quantity - Unit - Income Balance - Outcome Balance - Income Quantity 
    - Outcome Quantity - Current Quantity - Quantity Of Start) 
    [Detailed Items] (
        Name - Code - Unit - Start Quantity
        Table => [ Income Quantity - Outcome Quantity - Current Quantity of this transaction - Income Balance 
        - Outcome Balance - Client - Transaction Type]
    )
    """
    # Quick
    data = {}
    data['items'] = []
    ctx = {
        **companyDetails()
    }

    # Get Params
    kind = request.GET.get('kind', None)
    data['fromItem'] = fromItem = request.GET.get('fromItem', None)
    data['toItem'] = toItem = request.GET.get('toItem', None)
    fromDate = request.GET.get('fromDate')
    toDate = request.GET.get('toDate')

    if (not kind) :
        return render(request, 'finance/reportPage.html', ctx)
    elif kind == 'general' :
        template = 'finance/itemReportGeneral.html'
    elif kind == 'details' :
        template = 'finance/itemReportDetails.html'

    # If both Params of any field is empty return Error
    if (not fromItem and not toItem) :
        ctx['message'] = 'برجاء ادخال الاصناف '
        return render(request, 'finance/reportPage.html', ctx)

    # Check if None then use second value for item
    if fromItem :
        fromItem = Items.objects.get(Q(code=fromItem) | Q(item=fromItem)).pk
    else :
        fromItem = Items.objects.first().pk

    if toItem :
        toItem = Items.objects.get(Q(code=toItem) | Q(item=toItem)).pk
    else :
        toItem = Items.objects.last().pk

    # Prepare Query Filter
    kwrags = {}

    # Make Time filter
    time = {}
    if fromDate:
        fromDate = datetime.datetime.strptime(fromDate, '%Y-%m-%d')
        time['transaction__dateTime__gte'] = fromDate
    else :
        data['fromDate'] = Transactions.objects.first().getDate()

    if toDate:
        toDate += ' 23:59:59'
        toDate = datetime.datetime.strptime(toDate, '%Y-%m-%d %H:%M:%S')
        time['transaction__dateTime__lte'] = toDate
    else :
        data['toDate'] = Transactions.objects.last().getDate()

    # Make pk in order for filtering
    if fromItem < toItem :
        kwrags['pk__gte'] = fromItem
        kwrags['pk__lte'] = toItem
    elif fromItem > toItem :
        kwrags['pk__lte'] = fromItem
        kwrags['pk__gte'] = toItem
    else :
        kwrags['pk'] = fromItem


    if kind == 'general' :
        fromInventory = request.GET.get('fromInventory', None)
        toInventory = request.GET.get('toInventory', None)
        kwragsInv = {}
  
        # Check if there is inventory value
        if fromInventory or toInventory:
            if fromInventory :
                try :    
                    fromInventory = Inventories.objects.get(Q(inventory=fromInventory) | Q(pk=fromInventory))
                except :    
                    fromInventory = Inventories.objects.get(Q(inventory=fromInventory))
            else :
                fromInventory = Inventories.objects.first()
            data['fromInventory'] = fromInventory.inventory
            fromInventoryPK = fromInventory.pk
            

            if toInventory :
                data['toInventory'] = toInventory
                try :
                    toInventory =  Inventories.objects.get(Q(inventory=toInventory) | Q(pk=toInventory))
                except :    
                    toInventory =  Inventories.objects.get(Q(inventory=toInventory))
            else :
                toInventory = Inventories.objects.last()
            data['toInventory'] = toInventory.inventory
            toInventoryPK = toInventory.pk

            if fromInventoryPK < toInventoryPK :
                kwragsInv['inventory__pk__gte'] = fromInventoryPK
                kwragsInv['inventory__pk__lte'] = toInventoryPK
            elif fromInventoryPK > toInventoryPK :
                kwragsInv['inventory__pk__lte'] = fromInventoryPK
                kwragsInv['inventory__pk__gte'] = toInventoryPK
            else :
                kwragsInv['inventory__pk'] = fromInventoryPK

            itemSet = [ item['item'] for item in Inventory_Items.objects.filter(**kwragsInv).distinct().values('item')]
            kwrags['pk__in'] = itemSet
        else :
            fromInventory = Inventories.objects.first()
            data['fromInventory'] = fromInventory.inventory
            toInventory = Inventories.objects.last()
            data['toInventory'] = toInventory.inventory

        items = Items.objects.filter(**kwrags).order_by('pk')
        if not data['fromItem']:
            data['fromItem'] = items[0].item
        if not data['toItem']:
            data['toItem'] = items[len(items) - 1].item
        # (Name - Code - Quantity - Unit - Income Balance - Outcome Balance - Income Quantity - Outcome Quantity - Current Quantity - Quantity Of Start) 
        for item in items:
            kwrags = {
                'item' : item,
                'transaction__transactionType' : 'purchase',
                **kwragsInv,
                **time
            }
            name  = item.item
            code  = item.code
            unit  = item.unit
            price = item.sellPrice

            # Income Quantity
            purchaseTransactions = Transactions_Items.objects.filter(**kwrags).order_by('pk')
            incomeQuantity = purchaseTransactions.aggregate(Sum('quantity'))['quantity__sum']
            incomeBalance = math.trunc(sum([transactionPrice.getTotalPrice() for transactionPrice in purchaseTransactions]) * 100) / 100

            # Outcome Quantity
            kwrags['transaction__transactionType'] = 'sell'
            sellTransactions = Transactions_Items.objects.filter(**kwrags).order_by('pk')
            outcomeQuantity = sellTransactions.aggregate(Sum('quantity'))['quantity__sum'] or 0
            outcomeBalance = math.trunc(sum([transactionPrice.getTotalPrice() for transactionPrice in sellTransactions]) * 100) / 100

            # Quantity of Start
            quantityOfStart = sum([itemQuantity.quantityFromLastYear for itemQuantity in item.inventoriesList.filter(**kwragsInv)])
            quantity = sum([itemQuantity.quantity for itemQuantity in item.inventoriesList.filter(**kwragsInv)])
            totalPrice = quantity * math.trunc(price * 100) / 100
            # Add Data 
            data['items'].append({
                'item' : name,
                'code' : code,
                'currentQuantity' : quantity,
                'unit' : unit,
                'price' : price,
                'totalBalance' : totalPrice,
                'incomeQuantity' : incomeQuantity,
                'incomeBalance' : incomeBalance,
                'outcomeQuantity' : outcomeQuantity,
                'outcomeBalance' : outcomeBalance,
                'quantityOfStart' : quantityOfStart,
            })
        # print(data)

    elif kind == 'details':
        items = Items.objects.filter(**kwrags).order_by('pk')
        # Name - Code - Unit - Start Quantity
        # Table => [ Income Quantity - Outcome Quantity - Current Quantity of this transaction - Income Balance 
        # - Outcome Balance - Client - Transaction Type]

        for item in items :
            name  = item.item
            code  = item.code
            unit  = item.unit
            quantity  = item.quantity
            currentPerTransaction = quantityOfStart = sum([q.quantityFromLastYear for q in item.inventoriesList.all()]) or 0

            # Table
            totalIncomingQuantity = 0
            totalOutcomingQuantity = 0
            totalIncomingBalance = 0
            totalOutcomingBalance = 0
            table = []
            for transaction in item.transactionsList.all().order_by('pk'):
                price = math.trunc(transaction.price * 100) / 100
                quantityPerTransaction = transaction.quantity
                inventory = transaction.inventory.inventory if transaction.inventory else None
                client = transaction.transaction.user
                transactionType =  transaction.transaction.transactionType
                total = math.trunc(price*100) * quantityPerTransaction / 100

                if transactionType == 'sell':
                    transactionType = 'فاتوره مبيعات'
                    currentPerTransaction -= quantityPerTransaction
                    totalOutcomingQuantity += quantityPerTransaction
                    totalIncomingBalance += total
                    incomeQuantity = 0
                    outcomeQuantity = quantityPerTransaction
                    
                elif transactionType == 'purchase':
                    transactionType = 'فاتوره مشتريات'
                    currentPerTransaction += quantityPerTransaction
                    totalIncomingQuantity += quantityPerTransaction
                    totalOutcomingBalance += total
                    incomeQuantity = quantityPerTransaction
                    outcomeQuantity = 0
                
                
                if transaction in item.transactionsList.filter(**time):
                    table.append({
                        'price' : price,
                        'incomeQuantity' : incomeQuantity,
                        'outcomeQuantity' : outcomeQuantity,
                        'currentQuantity' : currentPerTransaction,
                        'transactionType' : transactionType,
                        'inventory' : inventory,
                        'client' : client,
                        'total' : total,
                    })
            
            data['items'].append({
                'item' : name,
                'code' : code,
                'unit' : unit,
                'quantity' : quantity,
                'quantityOfStart' : quantityOfStart,
                'totalIncomingQuantity' : totalIncomingQuantity,
                'totalOutcomingQuantity' : totalOutcomingQuantity,
                'totalIncomingBalance' : totalIncomingBalance,
                'totalOutcomingBalance' : totalOutcomingBalance,
                'table' : table,
                'profit' : math.trunc((totalIncomingBalance - totalOutcomingBalance) * 100) / 100,
            })
        
    else :
        return JsonResponse({'message' : 'Invalid Route!'}, status=404)
        
    return render(request, template, {
        'data' : data,
        **companyDetails()
    })

@login_required(login_url="loginView")
def profile(request):
    if request.method == 'GET' :
        message = request.session.get('message', None)
        success = request.session.get('success', None)
        if message :
            del request.session['message']
            del request.session['success']
        user = User.objects.get(username=request.user.username)
        username = user.username
        firstName = user.first_name
        lastName = user.last_name
        settingsData = settings(request)
        return render(request, 'finance/profile.html', {
            'username' : username,
            'name' : f'{firstName} {lastName}',
            'settings' : settingsData,
            'message' : message,
            'success' : success,
            'arName' : companyDetails()['arName']
        })
    elif request.method == 'POST' :
        username = request.POST.get('username', '')
        try :
            firstName, lastName = request.POST.get('name', '').split()
        except :
            firstName = request.POST.get('name', '')
            lastName = ''
        password = request.POST.get('password', None)
        if not (username and  firstName) :
            request.session['message'] = 'برجاء ادخال اسم مستخدم و الاسم الاول'
            request.session['success'] = 'false'
        else :
            user = User.objects.filter(username=request.user.username)
            if user and username != request.user.username:
                request.session['message'] = 'اسم المستخدم موجود من قبل برجاء استخدام غيره'
                request.session['success'] = 'false'

            else :
                user = User.objects.get(username=request.user.username)
                user.first_name = firstName
                user.last_name = lastName
                if password :
                    user.set_password(password)
                user.save()
                request.session['message'] = 'تم تعديل البيانات بنجاح!'
                request.session['success'] = 'true'


        return HttpResponseRedirect(reverse('profile'))
    else :
        return JsonResponse({'message' : 'Invalid Route'}, status=404)
    
def settings(request):
    data = Settings.objects.first()

    if request.method == 'GET' :
        if data :
            data = data.serialize()
        else :
            data = {}
        return data
    
    elif request.method == 'POST' :
        companyEnglishName = request.POST.get('engName', None)
        companyArabicName = request.POST.get('arName', None)
        taxFileNumber = request.POST.get('taxFileNumber', None)
        taxRegistrationNumber = request.POST.get('taxRegistrationNumber', None)
        tax3Number = request.POST.get('tax3Number', None)
        phoneNumber = request.POST.get('phoneNumber', None)
        address = request.POST.get('address', None)
        if taxFileNumber and taxRegistrationNumber and tax3Number :
            request.session['message'] = 'برجاء ادخال البيانات الضريبيه'
            request.session['success'] = 'false'

        if data :
            data.companyEnglishName =  companyEnglishName
            data.companyArabicName = companyArabicName
            data.taxRegistrationNumber = taxFileNumber
            data.taxFileNumber = taxRegistrationNumber
            data.tax3Number = tax3Number
            data.phoneNumber = phoneNumber
            data.address = address
            try :
                data.save()
                request.session['message'] = 'تم تعديل البيانات بنجاح!'
                request.session['success'] = 'true'


            except :
                request.session['message'] = "برجاء المحاوله مره اخري!"
                request.session['success'] = 'false'

        else :
            try :
                Settings.objects.create(
                    companyEnglishName =  companyEnglishName,
                    companyArabicName = companyArabicName,
                    taxRegistrationNumber = taxFileNumber,
                    taxFileNumber = taxRegistrationNumber,
                    tax3Number = tax3Number,
                    phoneNumber = phoneNumber,
                    address = address
                )
                request.session['message'] = 'تم تسجيل البيانات بنجاح!'
                request.session['success'] = 'true'

            except Exception as e :
                print(str(e))
                request.session['message'] = "برجاء المحاوله مره اخري!"
                request.session['success'] = 'false'

        return HttpResponseRedirect(reverse('profile'))
    else :
        return JsonResponse({'message' : 'Invalid Route'}, status=404)


import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL as WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from django.http import StreamingHttpResponse, FileResponse
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
import aspose.words as aw


import io

def generate_word(request):
    data = Settings.objects.first().serialize()
    document = docx.Document()
    section = document.sections[0]

    # Setting document for arabic style
    # font = document.font
    # font.complex_script = True
    # font.rtl = True
    # document.styles['Arabic'].font.language = 'ar'

    # Header 
    header = section.header
    
    # Header Table
    t1 = document.add_table(rows=4, cols=3)
    tblPr = t1._element.xpath('.//w:tblPr')[0]

    
    bidiVisual = OxmlElement('w:bidiVisual')
    tblPr.append(bidiVisual)

    # Add Data
    row1 = t1.rows[0].cells
    row1[0].text = data['arName']
    row1[2].text = f'رقم الملف الضريبي : {data["taxFileNumber"]}'

    row2 = t1.rows[1].cells
    row2[0].text = data['engName']
    row2[2].text = f"رقم التسجيل الضريبي : {data['taxRegistrationNumber']}"

    row3 = t1.rows[2].cells
    row3[0].text = f"العنوان : {data['address']}"
    row3[2].text = f"البطاقه الضريبيه : {data['tax3Number']}"

    row4 = t1.rows[3].cells
    row4[0].text = f"تليفون : {data['phoneNumber']}"

    # Remove Border
    for row in t1.rows:
        for cell in row.cells:
            tc = cell._element.tcPr
            tc.left = None
            tc.top = None
            tc.right = None
            tc.bottom = None
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    

    header.table = t1
    # send response
    # save document info
    # buffer = io.BytesIO()
    document.save('thisisdoc.docx')  # save your memory stream
    # buffer.seek(0)  # rewind the stream
    doc = aw.Document("thisisdoc.docx")

    # Save as PDF
    doc.save("PDF.pdf")

    response = HttpResponse("PDF.pdf", content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="'"PDF.pdf"'"'

    # put them to streaming content response 
    # within docx content_type
    # response = StreamingHttpResponse(
    #     streaming_content=buffer,  # use the stream's content
    #     content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    # )

    # response['Content-Disposition'] = 'attachment;filename=Test.docx'
    response["Content-Encoding"] = 'UTF-8'
    return response
    pass
